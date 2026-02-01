"""Document ingestion: load PDF/TXT/DOC, chunk, and prepare for indexing."""

from datetime import datetime, timezone
from pathlib import Path
from typing import List

from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.core.config import get_settings
from app.core.logger import get_logger

logger = get_logger(__name__)

# Lazy load for optional dependencies
def _load_docx(file_path: str) -> List[Document]:
    try:
        from langchain_community.document_loaders import UnstructuredWordDocumentLoader
        loader = UnstructuredWordDocumentLoader(file_path)
        return loader.load()
    except Exception:
        import docx
        doc = docx.Document(file_path)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        return [Document(page_content="\n\n".join(texts), metadata={"source": file_path})]


def load_documents(file_path: str) -> List[Document]:
    """
    Load a single file into LangChain Documents.
    Supports PDF, TXT, DOC, DOCX.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        loader = PyPDFLoader(file_path)
        docs = loader.load()
    elif suffix in (".txt", ".md"):
        loader = TextLoader(file_path, encoding="utf-8", autodetect_encoding=True)
        docs = loader.load()
    elif suffix in (".doc", ".docx"):
        docs = _load_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    return docs


def chunk_documents(
    documents: List[Document],
    filename: str,
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
) -> List[Document]:
    """
    Split documents into chunks with metadata: filename, page, upload_time.
    """
    settings = get_settings()
    size = chunk_size or settings.chunk_size
    overlap = chunk_overlap or settings.chunk_overlap
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=size,
        chunk_overlap=overlap,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    upload_time = datetime.now(timezone.utc).isoformat()
    chunks: List[Document] = []
    for doc in documents:
        page = doc.metadata.get("page", doc.metadata.get("source", ""))
        meta_base = {
            "filename": filename,
            "page": page,
            "upload_time": upload_time,
        }
        for chunk in splitter.split_documents([doc]):
            chunk.metadata = {**meta_base, **chunk.metadata}
            chunks.append(chunk)
    logger.info("Chunked %s into %d chunks", filename, len(chunks))
    return chunks


def ingest_file(file_path: str, filename: str) -> List[Document]:
    """
    Load and chunk a file; return list of Document chunks with metadata.
    """
    raw_docs = load_documents(file_path)
    return chunk_documents(raw_docs, filename)
